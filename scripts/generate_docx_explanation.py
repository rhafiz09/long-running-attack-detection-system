import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_styled_word_doc():
    doc = docx.Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Define color palette (Professional Deep Blue & Slate Gray)
    COLOR_PRIMARY = RGBColor(15, 32, 67)       # Deep Navy Blue (#0F2043)
    COLOR_SECONDARY = RGBColor(41, 128, 185)   # Accent Blue (#2980B9)
    COLOR_TEXT = RGBColor(44, 62, 80)          # Dark Slate Gray (#2C3E50)
    COLOR_MUTED = RGBColor(127, 140, 141)      # Muted Gray (#7F8C8D)
    COLOR_CODE_BG = RGBColor(240, 243, 246)    # Code Block Light Gray Blue
    COLOR_CODE_TXT = RGBColor(199, 37, 78)     # Code Crimson / Dark Crimson
    HEX_BG_LIGHT = "F4F6F9"                    # Light Gray-Blue for tables/callouts
    HEX_CODE_BG = "F0F3F6"
    HEX_PRIMARY = "0F2043"
    HEX_ACCENT = "2980B9"

    # Style definitions
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = COLOR_TEXT
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    def set_cell_background(cell, hex_color):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{margin_name}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def add_code_block(text_code):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_background(cell, HEX_CODE_BG)
        set_cell_margins(cell, 120, 120, 150, 150)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text_code)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(35, 43, 43)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def add_header_footer():
        section = doc.sections[0]
        header = section.header
        hp = header.paragraphs[0]
        hp.text = "ARSITEKTUR & INTEGRASI SISTEM DETEKSI SERANGAN LONG-RUNNING (CNN-LSTM)"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.style.font.size = Pt(8.5)
        hp.style.font.color.rgb = COLOR_MUTED

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = "Dokumen Arsitektur Teknis Sistem • Pemetaan Referensi Baris Kode (Codebase Mapping)"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.style.font.size = Pt(8.5)
        fp.style.font.color.rgb = COLOR_MUTED

    add_header_footer()

    # --- TITLE ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("DOKUMEN PENJELASAN TEKNIS & ARSITEKTURAL")
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_PRIMARY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run("Integrasi Antar-Komponen & Referensi Baris Kode dalam Sistem Deteksi\nSerangan Siber Long-Running Berbasis Deep Learning Hybrid CNN-LSTM")
    run_sub.font.size = Pt(12)
    run_sub.font.bold = True
    run_sub.font.color.rgb = COLOR_SECONDARY

    # --- RINGKASAN EKSEKUTIF ---
    h1 = doc.add_heading(level=1)
    run_h1 = h1.add_run("1. RINGKASAN EKSEKUTIF & TINJAUAN ARSITEKTUR")
    run_h1.font.color.rgb = COLOR_PRIMARY
    run_h1.font.size = Pt(14)
    run_h1.font.bold = True

    doc.add_paragraph(
        "Sistem Pendeteksi Serangan Long-Running (Long-Running Attack Detection System) dirancang sebagai platform "
        "pertahanan keamanan siber full-stack berskala Enterprise. Ancaman persisten tingkat lanjut (Advanced Persistent Threats / APT) "
        "modern kerap berhasil menghindari deteksi aturan firewall konvensional yang bersifat statis dengan cara memperlambat "
        "dan memperpanjang durasi penyerangan (serangan bertipe low-and-slow). Aktivitas berbahaya seperti Internal Reconnaissance, "
        "Lateral Movement, dan Beaconing dapat berlangsung berjam-jam, berhari-hari, hingga berminggu-minggu."
    )
    doc.add_paragraph(
        "Untuk mengatasi tantangan kompleks ini, sistem dibangun di atas arsitektur terdekopel (Decoupled Layered Microservices Architecture) "
        "yang memisahkan secara tegas antara lapisan presentasi (Frontend SOC Dashboard), lapisan pemrosesan dan manajemen log "
        "(Backend API Engine), lapisan penyimpanan data terpusat (Relational Database), lapisan analisis kecerdasan buatan (AI Deep Learning Engine), "
        "serta lapisan pemrosesan latar belakang (Asynchronous Pipeline Worker & Message Queuing). Dokumen ini menguraikan secara komprehensif "
        "empat pilar integrasi utama beserta referensi eksplisit baris kode (codebase mapping) pada proyek nyata Anda."
    )

    # --- DIAGRAM ARSITEKTUR TABEL ---
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    table_arch = doc.add_table(rows=1, cols=1)
    table_arch.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_arch = table_arch.cell(0, 0)
    set_cell_background(cell_arch, HEX_BG_LIGHT)
    set_cell_margins(cell_arch, 150, 150, 200, 200)
    p_box = cell_arch.paragraphs[0]
    p_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_box = p_box.add_run(
        "DIAGRAM ALIR INTEGRASI EMPAT PILAR KOMPONEN SISTEM\n\n"
        "[ Frontend / Kolektor / Browser ]  <=== (1) REST API berformat JSON ===>  [ Backend FastAPI & Django ]\n"
        "                                                                                    |        |\n"
        "               (3) Direct Python Integration                                        |        | (2) Django ORM\n"
        "                   (Tensor 3D / Keras)                                              |        |    & SQLAlchemy\n"
        "                            v                                                       v        v\n"
        "[ Modul AI Hybrid CNN-LSTM (Inference Service) ]                      [ Basis Data PostgreSQL 15 (ORM) ]\n"
        "                            ^                                                        ^\n"
        "                            |======= (4) Asynchronous Operations / Queue Worker =====|"
    )
    run_box.font.size = Pt(9.5)
    run_box.font.bold = True
    run_box.font.color.rgb = COLOR_PRIMARY
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- 2. FRONTEND KE BACKEND ---
    h2 = doc.add_heading(level=1)
    run_h2 = h2.add_run("2. FRONTEND KE BACKEND: MENGGUNAKAN REST API DALAM FORMAT JSON")
    run_h2.font.color.rgb = COLOR_PRIMARY
    run_h2.font.size = Pt(14)
    run_h2.font.bold = True

    doc.add_paragraph(
        "Integrasi antara antarmuka pengguna (Frontend SOC Dashboard berbasis Django/HTML5 dengan telemetri Chart.js & Tailwind CSS) "
        "maupun perangkat kolektor eksternal dengan server Backend (FastAPI di Port 8000 dan Django Backend di Port 8001) sepenuhnya "
        "dijembatani oleh protokol HTTP/HTTPS menggunakan gaya arsitektur REST API (Representational State Transfer) dengan struktur "
        "payload berbasis JSON (JavaScript Object Notation)."
    )

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r1 = p.add_run("a. Standarisasi Format Pertukaran Data (JSON Schema & Pydantic Validation):\n")
    r1.font.bold = True
    r1.font.color.rgb = COLOR_SECONDARY
    p.add_run(
        "Seluruh permintaan (Request) dan tanggapan (Response) distandarisasi menggunakan format JSON yang ringan dan universal. "
        "Di sisi Backend FastAPI, setiap payload JSON yang masuk divalidasi menggunakan skema Pydantic (`schemas/detection.py`). "
        "Validasi ini menjamin tipe data yang masuk (seperti IP Origin, IP Impacted, Port Impacted, Zone, dan Log Date) sesuai "
        "dengan skema 112+ kolom sebelum diproses lebih lanjut."
    )

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r2 = p.add_run("b. Mekanisme Komunikasi Data Real-Time & Asinkron:\n")
    r2.font.bold = True
    r2.font.color.rgb = COLOR_SECONDARY
    p.add_run(
        "• Telemetri Pemantauan Real-Time (Chart.js): UI melakukan request JSON berkala ke endpoint `/api/chart-data/` atau `/api/v1/detect` "
        "untuk memperbarui tren grafik serangan dan statistik tabel log (Palo Alto, FortiGate, FortiWAF).\n"
        "• Integrasi Chatbot AI Assistant SOC (Google Gemini): Analis SOC mengirimkan log melalui antarmuka chat. Frontend mengirim payload JSON "
        "berisi teks log ke backend, yang kemudian meneruskannya ke engine LLM dan mengembalikan response JSON terstruktur berisi analisis sasaran IP, "
        "jenis serangan (DDoS/Recon/Lateral Movement), kecepatan serangan, dan rekomendasi mitigasi."
    )

    p_ref1 = doc.add_paragraph()
    run_ref1 = p_ref1.add_run("📌 REFERENSI BARIS KODE (CODEBASE MAPPING):")
    run_ref1.font.bold = True
    run_ref1.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph(
        "1. File `app/api/routers/detection.py` (Baris 13-42): Endpoint POST /detect untuk inferensi real-time menerima payload JSON `DetectionRequest`, "
        "melakukan pengecekan `X-API-Key`, memanggil `inference_service.predict(payload)`, dan mengembalikan `DetectionResponse` berbasis JSON:"
    )
    add_code_block(
        "@router.post(\"\", response_model=DetectionResponse, status_code=status.HTTP_200_OK,\n"
        "             dependencies=[Depends(verify_api_key)])\n"
        "@limiter.limit(\"100/minute\")\n"
        "def detect_attacks(request: Request, payload: DetectionRequest,\n"
        "                   inference_service: InferenceService = Depends(get_inference_service)) -> DetectionResponse:\n"
        "    logger.debug(f\"POST /api/v1/detect requested with {len(payload.logs)} log records.\")\n"
        "    return inference_service.predict(payload)"
    )

    doc.add_paragraph(
        "2. File `web_dashboard/monitor/views.py` (Baris 319-408): Endpoint REST API Django `chart_data_api(request)` merespons permintaan polling "
        "grafik Chart.js dari browser dan mengembalikan struktur `JsonResponse` berisi agregasi lalu lintas berdasarkan filter waktu ('30s', '1m', '24h'):"
    )
    add_code_block(
        "@login_required(login_url=\"/login/\")\n"
        "def chart_data_api(request):\n"
        "    time_filter = request.GET.get(\"filter\", \"24h\")\n"
        "    # ... (Kalkulasi agregasi log dari PaloAltoLog, FortinetLog, FortiwafLog)\n"
        "    return JsonResponse({\n"
        "        \"status\": \"success\", \"filter\": time_filter, \"labels\": labels,\n"
        "        \"datasets\": { \"palo_alto\": pa_data, \"fortinet\": fn_data, \"fortiwaf\": fw_data }\n"
        "    })"
    )

    doc.add_paragraph(
        "3. File `web_dashboard/monitor/views.py` (Baris 410-484): Endpoint `chatbot_api(request)` mengurai payload JSON dari pesan analis SOC, "
        "mengekstrak IP, meminta analisis ke model Google Gemini (`gemini-2.5-flash`), dan mengembalikan `JsonResponse` terstruktur:"
    )
    add_code_block(
        "@csrf_exempt\n"
        "@require_POST\n"
        "@login_required(login_url=\"/login/\")\n"
        "def chatbot_api(request):\n"
        "    body = json.loads(request.body.decode(\"utf-8\"))\n"
        "    user_message = body.get(\"message\", \"\").strip()\n"
        "    # ... (Ekstraksi IP & integrasi Gemini AI / Simulation Fallback)\n"
        "    return JsonResponse({\"status\": \"success\", \"response\": response.text, \"target_ip\": target_ip})"
    )

    # --- 3. BACKEND KE BASIS DATA ---
    h3 = doc.add_heading(level=1)
    run_h3 = h3.add_run("3. BACKEND KE BASIS DATA: MENGGUNAKAN DJANGO ORM & SQLALCHEMY")
    run_h3.font.color.rgb = COLOR_PRIMARY
    run_h3.font.size = Pt(14)
    run_h3.font.bold = True

    doc.add_paragraph(
        "Sistem menggunakan PostgreSQL 15 sebagai basis data relasional utama (`RDBMS`) untuk menampung jutaan log aktivitas jaringan dan "
        "hasil klasifikasi AI. Untuk menjamin keandalan, keamanan, serta pemisahan tanggung jawab (Separation of Concerns), interaksi antara "
        "Backend dan Basis Data diimplementasikan menggunakan teknik Object-Relational Mapping (ORM)."
    )

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r_db1 = p.add_run("a. Arsitektur ORM Terdekopel (Unmanaged Django ORM + Managed SQLAlchemy):\n")
    r_db1.font.bold = True
    r_db1.font.color.rgb = COLOR_SECONDARY
    p.add_run(
        "• Django ORM (`web_dashboard/monitor/models.py`): Digunakan oleh aplikasi Django Dashboard untuk pembacaan data, paginasi, pemfilteran, "
        "dan rendering telemetri UI. Model ORM dideklarasikan dengan parameter `Meta.managed = False`. Hal ini memastikan bahwa Django ORM beroperasi "
        "dalam mode Read-Only / Unmanaged sehingga tidak akan melakukan migrasi atau mengubah skema fisik tabel yang sedang menampung aliran data berkecepatan tinggi.\n"
        "• SQLAlchemy ORM (`app/models/logs.py`): Digunakan oleh FastAPI ML Engine dan Pipeline Worker untuk melakukan penciptaan tabel, "
        "manajemen transaksi, dan High-Throughput Batch Insertion."
    )

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r_db2 = p.add_run("b. Manajemen Skema Dinamis dengan PostgreSQL JSONB (`SafeJSONField`):\n")
    r_db2.font.bold = True
    r_db2.font.color.rgb = COLOR_SECONDARY
    p.add_run(
        "Setiap vendor firewall memiliki atribut log khusus di luar kolom standar. Untuk mengakomodasi 112+ kolom atribut tanpa menyebabkan "
        "pembengkakan tabel, Django ORM memanfaatkan `SafeJSONField` yang dipetakan langsung ke tipe data native `JSONB` pada PostgreSQL. "
        "ORM melakukan serialisasi dan deserialisasi otomatis dari struktur JSON di database menjadi objek `dict/list` Python."
    )

    p_ref2 = doc.add_paragraph()
    run_ref2 = p_ref2.add_run("📌 REFERENSI BARIS KODE (CODEBASE MAPPING):")
    run_ref2.font.bold = True
    run_ref2.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph(
        "1. File `web_dashboard/monitor/models.py` (Baris 4-18 & 50-60): Implementasi `SafeJSONField` dan model `PaloAltoLog` dengan atribut "
        "`managed = False` dan pemetaan tabel `db_table = \"palo_alto_logs\"`:"
    )
    add_code_block(
        "class SafeJSONField(models.JSONField):\n"
        "    def from_db_value(self, value, expression, connection):\n"
        "        if value is None or isinstance(value, (dict, list, int, float, bool)):\n"
        "            return value\n"
        "        try: return super().from_db_value(value, expression, connection)\n"
        "        except TypeError: return value\n\n"
        "class PaloAltoLog(BaseFirewallLog):\n"
        "    class Meta:\n"
        "        managed = False  # Django beroperasi dalam mode Read-Only ORM\n"
        "        db_table = \"palo_alto_logs\"\n"
        "        ordering = [\"-log_date\"]"
    )

    doc.add_paragraph(
        "2. File `app/models/logs.py` (Baris 6-42 & 44-59): Model SQLAlchemy ORM yang mendefinisikan 7 kolom inti berindeks (`log_date`, "
        "`ip_origin`, `ip_impacted`, dll) dan kolom dinamis `additional_data = Column(JSONB)` untuk sisa 112+ kolom CSV:"
    )
    add_code_block(
        "class BaseLogModel(Base):\n"
        "    __abstract__ = True\n"
        "    id = Column(Integer, primary_key=True, index=True, autoincrement=True)\n"
        "    log_date = Column(DateTime(timezone=True), index=True, nullable=False)\n"
        "    ip_origin = Column(String(100), index=True, nullable=True)\n"
        "    # ... (Kolom indeks penting lainnya untuk inferensi ML & SOC)\n"
        "    additional_data = Column(JSONB, nullable=False, default=dict)\n\n"
        "class PaloAltoLog(BaseLogModel):\n"
        "    __tablename__ = \"palo_alto_logs\""
    )

    # --- 4. BACKEND KE MODUL ML ---
    h4 = doc.add_heading(level=1)
    run_h4 = h4.add_run("4. BACKEND KE MODUL PEMBELAJARAN MESIN: INTEGRASI BERBASIS PYTHON")
    run_h4.font.color.rgb = COLOR_PRIMARY
    run_h4.font.size = Pt(14)
    run_h4.font.bold = True

    doc.add_paragraph(
        "Kelebihan utama membangun Backend (FastAPI/Django) dan Modul Pembelajaran Mesin (TensorFlow/Keras/Scikit-learn) dalam satu ekosistem "
        "bahasa pemrograman Python adalah terciptanya Direct In-Memory Python Integration. Tidak ada latensi jaringan internal akibat panggilan "
        "REST API tambahan atau overhead serialisasi jaringan saat model AI melakukan klasifikasi log jaringan."
    )

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r_ml1 = p.add_run("a. Pola Arsitektur Singleton untuk Warm Model Loading (`InferenceService`):\n")
    r_ml1.font.bold = True
    r_ml1.font.color.rgb = COLOR_SECONDARY
    p.add_run(
        "Untuk menghindari keterlambatan (overhead) akibat memuat model berukuran besar berulang-ulang, sistem menggunakan pola desain Singleton "
        "(`app/services/inference_service.py`). Saat FastAPI atau Worker pertama kali berjalan, model neural network Keras (`.keras`) dan "
        "`LogFeatureEngineer` (`.pkl`) dimuat ke dalam RAM secara persisten. Model siap sedia (*warm instance*) untuk mengeksekusi inferensi "
        "dalam hitungan milidetik saat aliran log masuk."
    )

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r_ml2 = p.add_run("b. Preprocessing Langsung & Incremental Windowing (`LogFeatureEngineer`):\n")
    r_ml2.font.bold = True
    r_ml2.font.color.rgb = COLOR_SECONDARY
    p.add_run(
        "Sebelum tensor dikirim ke jaringan saraf, modul Python (`app/ai_engine/feature_engineering.py`) melakukan pemrosesan data secara langsung:\n"
        "• Incremental Windowing: Sistem menghitung jendela waktu beruntun sejak suatu IP pertama kali terlihat di jaringan hingga waktu evaluasi.\n"
        "• Feature Engineering Khusus Attack Stage: Ekstraksi frekuensi scanning port/host (Class 1: Internal Reconnaissance), pelacakan perpindahan IP "
        "internal-ke-internal berdasarkan riwayat host (Class 2: Lateral Movement), serta pendeteksian keteraturan pola interval waktu C2 (Class 3: Beaconing).\n"
        "• Evaluasi Model Hybrid CNN-LSTM: Data log diubah menjadi tensor Python 3-Dimensi `[samples, timesteps, features]`. Lapisan Conv1D mengekstrak pola "
        "spasial dan burst lokal, sedangkan lapisan LSTM mempertahankan memori jangka panjang."
    )

    p_ref3 = doc.add_paragraph()
    run_ref3 = p_ref3.add_run("📌 REFERENSI BARIS KODE (CODEBASE MAPPING):")
    run_ref3.font.bold = True
    run_ref3.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph(
        "1. File `app/services/inference_service.py` (Baris 188-196): Penggunaan dekorator `@lru_cache()` untuk menjamin pola Singleton "
        "sehingga pemuatan model `tf.keras.models.load_model()` hanya terjadi tepat 1 kali di awal startup aplikasi:"
    )
    add_code_block(
        "@lru_cache()\n"
        "def get_inference_service() -> InferenceService:\n"
        "    \"\"\"Singleton Dependency Provider: Caches the InferenceService instance using @lru_cache\"\"\"\n"
        "    return InferenceService()"
    )

    doc.add_paragraph(
        "2. File `app/services/inference_service.py` (Baris 42-58): Pemuatan langsung (*Direct In-Memory Loading*) artefak Keras "
        "`cnn_lstm_model.keras` dan objek scaler/encoder `feature_engineer.pkl` ke dalam RAM proses Python:"
    )
    add_code_block(
        "def _load_artifacts(self):\n"
        "    logger.info(f\"Loading Keras CNN-LSTM model from {MODEL_PATH}...\")\n"
        "    self.model = tf.keras.models.load_model(str(MODEL_PATH))\n"
        "    if ENGINEER_PATH.exists():\n"
        "        with open(ENGINEER_PATH, \"rb\") as f: self.engineer = pickle.load(f)\n"
        "        logger.info(\"LogFeatureEngineer loaded successfully.\")"
    )

    doc.add_paragraph(
        "3. File `app/services/inference_service.py` (Baris 111-137 & 155-178): Pemrosesan Preprocessing `clean_and_prepare_df()`, `perform_time_windowing()`, "
        "pembentukan Tensor 3D `[samples, timesteps, features]`, serta pemanggilan inferensi AI secara langsung dalam Python (`self.model.predict(X)`):"
    )
    add_code_block(
        "df_clean = engineer.clean_and_prepare_df(raw_logs)\n"
        "df_windowed = engineer.perform_time_windowing(df_clean)\n"
        "df_transformed = engineer.transform(df_labeled)\n\n"
        "# Pembentukan Tensor 3D berurutan dan pemanggilan model neural network\n"
        "X, ip_list = self.prepare_sequences_with_ips(engineer, df_transformed)\n"
        "predictions = self.model.predict(X, verbose=0)\n"
        "pred_labels = np.argmax(predictions, axis=1)  # Pemetaan Class 0 (Normal) hingga 3 (Beaconing)"
    )

    # --- 5. BACKEND KE ANTRIAN PESAN ---
    h5 = doc.add_heading(level=1)
    run_h5 = h5.add_run("5. BACKEND KE ANTRIAN PESAN: UNTUK OPERASI ASINKRON")
    run_h5.font.color.rgb = COLOR_PRIMARY
    run_h5.font.size = Pt(14)
    run_h5.font.bold = True

    doc.add_paragraph(
        "Dalam ekosistem Security Operations Center (SOC) skala enterprise, aliran data log dari banyak firewall dapat mencapai ribuan hingga jutaan "
        "event per detik. Jika proses pembersihan data, validasi 112+ kolom, penyimpanan ke database, dan evaluasi model AI CNN-LSTM dieksekusi secara "
        "sinkron (*synchronous blocking*) di dalam thread request API, maka server akan mengalami bottleneck, timeout, dan penurunan kinerja drastis. "
        "Oleh karena itu, sistem mengintegrasikan Operasi Asinkron dan Antrian Pesan (Asynchronous Task Processing & Background Workers)."
    )

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r_q1 = p.add_run("a. Background Pipeline Worker Asinkron (`pipeline_worker.py`):\n")
    r_q1.font.bold = True
    r_q1.font.color.rgb = COLOR_SECONDARY
    p.add_run(
        "Sistem menerapkan arsitektur pemrosesan terdekopel di mana `pipeline_worker.py` berjalan sebagai proses mandiri di latar belakang "
        "(daemon worker) yang terpisah dari server REST API tatap muka. Worker ini beroperasi dalam *continuous event loop* "
        "(interval 5 menit / 300 detik atau *on-demand trigger*) untuk mengambil batch data mentah dari sumber API kolektor (`siem-logs`), "
        "memotong dan mengklasifikasikan log berdasarkan jenis firewall, serta menyisipkannya ke tabel PostgreSQL secara *batch insertion*."
    )

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r_q2 = p.add_run("b. Post-Storage Event-Driven AI Triggering & Skalabilitas Enterprise (Celery + Redis):\n")
    r_q2.font.bold = True
    r_q2.font.color.rgb = COLOR_SECONDARY
    p.add_run(
        "• Post-Storage AI Triggering: Begitu data log baru selesai disisipkan ke database oleh pipeline worker, sistem memicu evaluasi inferensi AI "
        "di latar belakang tanpa memblokir koneksi API utama (*non-blocking processing*).\n"
        "• Skalabilitas Message Queue Broker: Untuk lingkungan berskala masif (jutaan log/menit), arsitektur ini siap diintegrasikan dengan broker antrian "
        "pesan seperti Redis/Celery atau RabbitMQ. Raw log yang masuk ke API segera dimasukkan ke dalam antrian broker (latensi < 5ms), dan sekumpulan "
        "*consumer worker* asinkron akan mengambil batch pesan untuk di-preprocess oleh model CNN-LSTM dengan jaminan *Zero Data Loss* dan *Retry Mechanism*."
    )

    p_ref4 = doc.add_paragraph()
    run_ref4 = p_ref4.add_run("📌 REFERENSI BARIS KODE (CODEBASE MAPPING):")
    run_ref4.font.bold = True
    run_ref4.font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph(
        "1. File `pipeline_worker.py` (Baris 281-297): Loop utama (*Continuous Event Loop*) yang mengorkestrasi operasi latar belakang asinkron "
        "setiap interval 5 menit (`interval=300s`) tanpa mengganggu server API tatap muka:"
    )
    add_code_block(
        "if args.loop:\n"
        "    logger.info(f\"Starting Continuous Pipeline Worker in loop mode (interval={args.interval}s)...\")\n"
        "    while True:\n"
        "        try:\n"
        "            current_offset = run_pipeline_tick(count=args.count, offset_range=current_offset)\n"
        "            logger.info(f\"Tick complete. Next offset cursor: {current_offset}. Sleeping for {args.interval}s...\")\n"
        "            time.sleep(args.interval)\n"
        "        except Exception as e:\n"
        "            logger.error(f\"Error in pipeline loop: {e}. Retrying in 15 seconds...\")\n"
        "            time.sleep(15)"
    )

    doc.add_paragraph(
        "2. File `pipeline_worker.py` (Baris 51-71): Mekanisme pelacakan posisi antrian (*Cursor Tracking*) `load_cursor()` dan `save_cursor()` "
        "yang menyimpan status offset di `.pipeline_cursor.json` agar antrian log tidak duplikat ataupun terputus saat terjadi restart:"
    )
    add_code_block(
        "def load_cursor() -> int:\n"
        "    if os.path.exists(CURSOR_FILE):\n"
        "        with open(CURSOR_FILE, \"r\") as f: return json.load(f).get(\"last_offset\", 0)\n"
        "    return 0\n\n"
        "def save_cursor(offset: int):\n"
        "    with open(CURSOR_FILE, \"w\") as f:\n"
        "        json.dump({\"last_offset\": offset, \"updated_at\": datetime.now(timezone.utc).isoformat()}, f)"
    )

    doc.add_paragraph(
        "3. File `pipeline_worker.py` (Baris 208-258): Eksekusi *Post-Storage Event-Driven Trigger* di mana setelah penyisipan massal ORM selesai, "
        "worker langsung memicu inferensi AI secara latar belakang dan mencatat peringatan ancaman (*Alerting*):"
    )
    add_code_block(
        "# Post-Storage Data Access & AI Triggering (Section 4.4.2.5)\n"
        "if inserted_total > 0 and len(active_ips) > 0:\n"
        "    logger.info(f\"Triggering Post-Storage AI CNN-LSTM Inference on {len(active_ips)} active IPs...\")\n"
        "    ai_payload = DetectionRequest(logs=detection_logs)\n"
        "    ai_response = inference_service.predict(ai_payload)\n"
        "    # ... (Penyimpanan ke tabel prediction_logs dan pencetakan log ancaman🚨)"
    )

    # --- 6. TABEL MATRIKS PERBANDINGAN INTEGRASI ---
    h6 = doc.add_heading(level=1)
    run_h6 = h6.add_run("6. MATRIKS PEMETAAN & PERBANDINGAN KOMPONEN INTEGRASI")
    run_h6.font.color.rgb = COLOR_PRIMARY
    run_h6.font.size = Pt(14)
    run_h6.font.bold = True

    headers = ["Jalur Integrasi", "Teknologi / Protokol", "Komponen & Baris Kode Terkait", "Fungsi & Keunggulan Utama"]
    rows_data = [
        [
            "Frontend ke Backend",
            "REST API\n(Format JSON)",
            "• app/api/routers/detection.py (L13-42)\n• app/api/routers/logs.py (L14-50)\n• web_dashboard/monitor/views.py (L319-408, L410-520)",
            "Pertukaran data berkinerja tinggi, divalidasi Pydantic Schema, telemetri Chart.js asinkron, serta integrasi AI Chatbot Gemini secara real-time."
        ],
        [
            "Backend ke Basis Data",
            "Django ORM &\nSQLAlchemy ORM",
            "• web_dashboard/monitor/models.py (L4-18, L50-60)\n• app/models/logs.py (L6-42, L44-59)",
            "Pemisahan tugas (Unmanaged Read-Only di Django & Managed Bulk Insert di SQLAlchemy), SafeJSONField dinamis, dan 100% bebas SQL Injection."
        ],
        [
            "Backend ke Modul ML",
            "Direct Python In-Memory\n(TensorFlow / Keras)",
            "• app/services/inference_service.py (L42-58, L111-137, L188-196)\n• app/ai_engine/feature_engineering.py",
            "Zero network latency, Singleton pattern untuk model warmup, konversi tensor 3D langsung untuk evaluasi CNN-LSTM 4-kelas (Recon, Lateral, Beaconing)."
        ],
        [
            "Backend ke Antrian Pesan",
            "Async Worker Loop,\nEvent-Driven / Celery+Redis",
            "• pipeline_worker.py (L51-71, L208-258, L281-297)\n• FUTURE_IMPROVEMENTS.md",
            "Pemrosesan non-blocking di latar belakang, shock absorber terhadap lonjakan log masif, retries otomatis, dan event-driven AI triggering."
        ]
    ]

    table_comp = doc.add_table(rows=len(rows_data)+1, cols=4)
    table_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_comp.style = 'Table Grid'

    # Format Header Row
    hdr_cells = table_comp.rows[0].cells
    for idx, header_text in enumerate(headers):
        hdr_cells[idx].text = header_text
        set_cell_background(hdr_cells[idx], HEX_PRIMARY)
        set_cell_margins(hdr_cells[idx], 120, 120, 150, 150)
        p_hdr = hdr_cells[idx].paragraphs[0]
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p_hdr.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)

    # Format Data Rows
    for r_idx, row_content in enumerate(rows_data):
        row_cells = table_comp.rows[r_idx+1].cells
        bg_hex = HEX_BG_LIGHT if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, cell_text in enumerate(row_content):
            row_cells[c_idx].text = cell_text
            set_cell_background(row_cells[c_idx], bg_hex)
            set_cell_margins(row_cells[c_idx], 100, 100, 120, 120)
            p_cell = row_cells[c_idx].paragraphs[0]
            p_cell.style.font.size = Pt(9.0)
            if c_idx == 0:
                p_cell.runs[0].font.bold = True
                p_cell.runs[0].font.color.rgb = COLOR_PRIMARY
            elif c_idx == 2:
                for run in p_cell.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- 7. KESIMPULAN ---
    h7 = doc.add_heading(level=1)
    run_h7 = h7.add_run("7. KESIMPULAN ARSITEKTUR & INTEGRASI KODE")
    run_h7.font.color.rgb = COLOR_PRIMARY
    run_h7.font.size = Pt(14)
    run_h7.font.bold = True

    doc.add_paragraph(
        "Integrasi keempat pilar teknologi beserta pemetaan kode sumber di atas membuktikan bahwa Sistem Pendeteksi Serangan "
        "Long-Running (CNN-LSTM) telah diimplementasikan dengan standar rekayasa perangkat lunak Enterprise (Clean Code, DRY, dan Layered Decoupling). "
        "Kecepatan pertukaran data via REST API JSON (`routers/detection.py`), keandalan manipulasi data menggunakan Django ORM (`monitor/models.py`), "
        "kedalaman analisis AI dengan integrasi langsung Python/TensorFlow (`services/inference_service.py`), serta ketahanan pemrosesan massal "
        "melalui antrian pesan asinkron (`pipeline_worker.py`) menciptakan harmoni kerja yang memungkinkan identifikasi serangan siber low-and-slow "
        "secara akurat, real-time, dan minim alarm palsu (False Positive/Miss Alarm)."
    )

    output_path = os.path.join(os.path.abspath("."), "Penjelasan_Integrasi_Sistem_Pendeteksi_Serangan_Long_Running.docx")
    doc.save(output_path)
    print(f"Document successfully created at: {output_path}")

if __name__ == "__main__":
    create_styled_word_doc()
